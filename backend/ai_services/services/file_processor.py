import os
import re
import logging
from typing import Dict, Any, Optional, List
from pathlib import Path
import tempfile

# PDF processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PDF processing libraries not available. Install with: pip install PyPDF2 pdfplumber")

# DOC/DOCX processing
try:
    from docx import Document
    import mammoth
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("DOCX processing libraries not available. Install with: pip install python-docx mammoth")

# Additional text processing
try:
    import chardet
    CHARDET_AVAILABLE = True
except ImportError:
    CHARDET_AVAILABLE = False
    logging.warning("chardet not available. Install with: pip install chardet")

class FileProcessor:
    """Process various file formats and extract meaningful text content."""
    
    def __init__(self):
        self.supported_formats = {
            '.txt': self._process_txt,
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_doc,
            '.rtf': self._process_rtf,
            '.md': self._process_markdown,
            '.html': self._process_html,
            '.htm': self._process_html
        }
    
    def process_file(self, file_path: str, file_content: Optional[bytes] = None) -> Dict[str, Any]:
        """
        Process a file and extract meaningful text content.
        
        Args:
            file_path: Path to the file or filename for content processing
            file_content: Raw file content as bytes (for uploaded files)
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            # Determine file format
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Process the file
            processor_func = self.supported_formats[file_extension]
            
            if file_content:
                # Process uploaded file content
                return processor_func(file_content, file_path)
            else:
                # Process file from path
                if not os.path.exists(file_path):
                    raise FileNotFoundError(f"File not found: {file_path}")
                return processor_func(file_path)
                
        except Exception as e:
            logging.error(f"Error processing file {file_path}: {str(e)}")
            raise
    
    def _process_txt(self, file_input: str, filename: str = None) -> Dict[str, Any]:
        """Process plain text files."""
        try:
            if isinstance(file_input, bytes):
                # Handle uploaded file content
                encoding = self._detect_encoding(file_input)
                try:
                    text = file_input.decode(encoding)
                except UnicodeDecodeError:
                    # Try with error handling
                    text = file_input.decode(encoding, errors='replace')
            else:
                # Handle file path
                encoding = self._detect_encoding_from_file(file_input)
                try:
                    with open(file_input, 'r', encoding=encoding) as f:
                        text = f.read()
                except UnicodeDecodeError:
                    # Try with error handling
                    with open(file_input, 'r', encoding=encoding, errors='replace') as f:
                        text = f.read()
            
            return {
                'text': self._clean_text(text),
                'metadata': {
                    'file_type': 'text',
                    'encoding': encoding,
                    'word_count': len(text.split()),
                    'character_count': len(text),
                    'line_count': len(text.splitlines())
                }
            }
        except Exception as e:
            logging.error(f"Error processing text file: {str(e)}")
            raise
    
    def _process_pdf(self, file_input: str, filename: str = None) -> Dict[str, Any]:
        """Process PDF files using multiple methods for better text extraction."""
        if not PDF_AVAILABLE:
            raise ImportError("PDF processing libraries not available")
        
        try:
            text_content = []
            metadata = {
                'file_type': 'pdf',
                'pages': 0,
                'word_count': 0,
                'character_count': 0
            }
            
            if isinstance(file_input, bytes):
                # Handle uploaded file content
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                    temp_file.write(file_input)
                    temp_file_path = temp_file.name
                
                try:
                    # Try pdfplumber first (better for complex layouts)
                    with pdfplumber.open(temp_file_path) as pdf:
                        metadata['pages'] = len(pdf.pages)
                        
                        for page_num, page in enumerate(pdf.pages):
                            page_text = page.extract_text()
                            if page_text:
                                text_content.append(f"Page {page_num + 1}:\n{page_text}\n")
                    
                    # If pdfplumber didn't extract much text, try PyPDF2
                    if len(''.join(text_content)) < 100:
                        with open(temp_file_path, 'rb') as file:
                            pdf_reader = PyPDF2.PdfReader(file)
                            metadata['pages'] = len(pdf_reader.pages)
                            
                            for page_num, page in enumerate(pdf_reader.pages):
                                page_text = page.extract_text()
                                if page_text:
                                    text_content.append(f"Page {page_num + 1}:\n{page_text}\n")
                
                finally:
                    os.unlink(temp_file_path)
            else:
                # Handle file path
                # Try pdfplumber first
                with pdfplumber.open(file_input) as pdf:
                    metadata['pages'] = len(pdf.pages)
                    
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(f"Page {page_num + 1}:\n{page_text}\n")
                
                # If pdfplumber didn't extract much text, try PyPDF2
                if len(''.join(text_content)) < 100:
                    with open(file_input, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        metadata['pages'] = len(pdf_reader.pages)
                        
                        for page_num, page in enumerate(pdf_reader.pages):
                            page_text = page.extract_text()
                            if page_text:
                                text_content.append(f"Page {page_num + 1}:\n{page_text}\n")
            
            extracted_text = '\n'.join(text_content)
            metadata['word_count'] = len(extracted_text.split())
            metadata['character_count'] = len(extracted_text)
            
            return {
                'text': self._clean_text(extracted_text),
                'metadata': metadata
            }
            
        except Exception as e:
            logging.error(f"Error processing PDF file: {str(e)}")
            raise
    
    def _process_docx(self, file_input: str, filename: str = None) -> Dict[str, Any]:
        """Process DOCX files."""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX processing libraries not available")
        
        try:
            if isinstance(file_input, bytes):
                # Handle uploaded file content
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                    temp_file.write(file_input)
                    temp_file_path = temp_file.name
                
                try:
                    doc = Document(temp_file_path)
                finally:
                    os.unlink(temp_file_path)
            else:
                # Handle file path
                doc = Document(file_input)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract text from tables
            tables = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(' | '.join(row_text))
                if table_text:
                    tables.append('\n'.join(table_text))
            
            # Combine all text
            all_text = '\n\n'.join(paragraphs + tables)
            
            return {
                'text': self._clean_text(all_text),
                'metadata': {
                    'file_type': 'docx',
                    'paragraphs': len(paragraphs),
                    'tables': len(tables),
                    'word_count': len(all_text.split()),
                    'character_count': len(all_text)
                }
            }
            
        except Exception as e:
            logging.error(f"Error processing DOCX file: {str(e)}")
            raise
    
    def _process_doc(self, file_input: str, filename: str = None) -> Dict[str, Any]:
        """Process DOC files using mammoth."""
        if not DOCX_AVAILABLE:
            raise ImportError("DOC processing libraries not available")
        
        try:
            if isinstance(file_input, bytes):
                # Handle uploaded file content
                with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                    temp_file.write(file_input)
                    temp_file_path = temp_file.name
                
                try:
                    with open(temp_file_path, 'rb') as file:
                        result = mammoth.extract_raw_text(file)
                finally:
                    os.unlink(temp_file_path)
            else:
                # Handle file path
                with open(file_input, 'rb') as file:
                    result = mammoth.extract_raw_text(file)
            
            extracted_text = result.value
            
            return {
                'text': self._clean_text(extracted_text),
                'metadata': {
                    'file_type': 'doc',
                    'word_count': len(extracted_text.split()),
                    'character_count': len(extracted_text),
                    'messages': result.messages
                }
            }
            
        except Exception as e:
            logging.error(f"Error processing DOC file: {str(e)}")
            raise
    
    def _process_rtf(self, file_input: str, filename: str = None) -> Dict[str, Any]:
        """Process RTF files."""
        try:
            if isinstance(file_input, bytes):
                # Handle uploaded file content
                encoding = self._detect_encoding(file_input)
                text = file_input.decode(encoding)
            else:
                # Handle file path
                encoding = self._detect_encoding_from_file(file_input)
                with open(file_input, 'r', encoding=encoding) as f:
                    text = f.read()
            
            # Basic RTF text extraction (remove RTF markup)
            cleaned_text = self._extract_rtf_text(text)
            
            return {
                'text': self._clean_text(cleaned_text),
                'metadata': {
                    'file_type': 'rtf',
                    'encoding': encoding,
                    'word_count': len(cleaned_text.split()),
                    'character_count': len(cleaned_text)
                }
            }
        except Exception as e:
            logging.error(f"Error processing RTF file: {str(e)}")
            raise
    
    def _process_markdown(self, file_input: str, filename: str = None) -> Dict[str, Any]:
        """Process Markdown files."""
        try:
            if isinstance(file_input, bytes):
                # Handle uploaded file content
                encoding = self._detect_encoding(file_input)
                text = file_input.decode(encoding)
            else:
                # Handle file path
                encoding = self._detect_encoding_from_file(file_input)
                with open(file_input, 'r', encoding=encoding) as f:
                    text = f.read()
            
            # Clean markdown formatting but keep structure
            cleaned_text = self._clean_markdown(text)
            
            return {
                'text': self._clean_text(cleaned_text),
                'metadata': {
                    'file_type': 'markdown',
                    'encoding': encoding,
                    'word_count': len(cleaned_text.split()),
                    'character_count': len(cleaned_text)
                }
            }
        except Exception as e:
            logging.error(f"Error processing Markdown file: {str(e)}")
            raise
    
    def _process_html(self, file_input: str, filename: str = None) -> Dict[str, Any]:
        """Process HTML files."""
        try:
            if isinstance(file_input, bytes):
                # Handle uploaded file content
                encoding = self._detect_encoding(file_input)
                text = file_input.decode(encoding)
            else:
                # Handle file path
                encoding = self._detect_encoding_from_file(file_input)
                with open(file_input, 'r', encoding=encoding) as f:
                    text = f.read()
            
            # Extract text from HTML
            cleaned_text = self._extract_html_text(text)
            
            return {
                'text': self._clean_text(cleaned_text),
                'metadata': {
                    'file_type': 'html',
                    'encoding': encoding,
                    'word_count': len(cleaned_text.split()),
                    'character_count': len(cleaned_text)
                }
            }
        except Exception as e:
            logging.error(f"Error processing HTML file: {str(e)}")
            raise
    
    def _detect_encoding(self, content: bytes) -> str:
        """Detect encoding of byte content."""
        if CHARDET_AVAILABLE:
            try:
                result = chardet.detect(content)
                encoding = result['encoding'] if result['encoding'] else 'utf-8'
                # Validate the encoding by trying to decode a small sample
                content[:100].decode(encoding)
                return encoding
            except (UnicodeDecodeError, LookupError):
                # Fallback to utf-8 with error handling
                return 'utf-8'
        return 'utf-8'
    
    def _detect_encoding_from_file(self, file_path: str) -> str:
        """Detect encoding of a file."""
        if CHARDET_AVAILABLE:
            try:
                with open(file_path, 'rb') as f:
                    raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding'] if result['encoding'] else 'utf-8'
                # Validate the encoding by trying to decode a small sample
                raw_data[:100].decode(encoding)
                return encoding
            except (UnicodeDecodeError, LookupError):
                # Fallback to utf-8 with error handling
                return 'utf-8'
        return 'utf-8'
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize quotes and dashes
        # Straight double and single quotes, and en/em dashes
        text = text.replace('“', '"').replace('”', '"')
        text = text.replace("‘", "'").replace("’", "'")
        text = text.replace('–', '-').replace('—', '-')
        
        return text.strip()
    
    def _extract_rtf_text(self, rtf_text: str) -> str:
        """Extract plain text from RTF content."""
        # Remove RTF control words and groups
        text = re.sub(r'\\[a-z]+\d*', '', rtf_text)
        text = re.sub(r'\{[^}]*\}', '', text)
        text = re.sub(r'\\[{}]', '', text)
        
        # Clean up remaining markup
        text = re.sub(r'\\\'[0-9a-fA-F]{2}', '', text)
        text = re.sub(r'\\\*\\[a-z]+', '', text)
        
        return text
    
    def _clean_markdown(self, markdown_text: str) -> str:
        """Clean markdown formatting while preserving structure."""
        # Remove markdown headers but keep text
        markdown_text = re.sub(r'^#{1,6}\s+', '', markdown_text, flags=re.MULTILINE)
        
        # Remove bold/italic markers
        markdown_text = re.sub(r'\*\*(.*?)\*\*', r'\1', markdown_text)
        markdown_text = re.sub(r'\*(.*?)\*', r'\1', markdown_text)
        markdown_text = re.sub(r'__(.*?)__', r'\1', markdown_text)
        markdown_text = re.sub(r'_(.*?)_', r'\1', markdown_text)
        
        # Remove code blocks but keep content
        markdown_text = re.sub(r'```.*?\n(.*?)```', r'\1', markdown_text, flags=re.DOTALL)
        markdown_text = re.sub(r'`(.*?)`', r'\1', markdown_text)
        
        # Remove links but keep text
        markdown_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', markdown_text)
        
        return markdown_text
    
    def _extract_html_text(self, html_text: str) -> str:
        """Extract plain text from HTML content."""
        # Remove HTML tags
        text = re.sub(r'<[^>]+>', '', html_text)
        
        # Decode HTML entities
        import html
        text = html.unescape(text)
        
        return text
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return list(self.supported_formats.keys())
    
    def is_supported(self, file_path: str) -> bool:
        """Check if file format is supported."""
        file_extension = Path(file_path).suffix.lower()
        return file_extension in self.supported_formats 